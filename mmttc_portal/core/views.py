from django.shortcuts import render, redirect

from django.contrib.auth import authenticate, login

from django.contrib.auth.models import User

from django.contrib.auth.decorators import login_required

from .models import Course
from .models import Applicant, Application
from .models import ContactMessage
from .models import EarlierAttendedCourse
from .models import CourseSchedule

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from django.core.files import File
from django.conf import settings
from datetime import timedelta
import os

from datetime import datetime

from docx import Document
from docx.shared import Inches
from django.utils import timezone
from django.core.files import File
import subprocess

def home(request):

    schedule = CourseSchedule.objects.first()

    return render(
        request,
        'public/index.html',
        {
            'schedule': schedule
        }
    )


# LOGIN

def login_page(request):

    if request.method == 'POST':

        username = request.POST.get('username')

        password = request.POST.get('password')

        user = authenticate(username=username,
                            password=password)

        if user is not None:

            login(request, user)

            if user.is_staff:

                return redirect('/admin-dashboard/')

            else:

                return redirect('/dashboard/')

    return render(request,
                  'public/login.html')


def register_page(request):

    if request.method == 'POST':

        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')

        if User.objects.filter(username=username).exists():

            return render(request,
                          'public/register.html',
                          {
                              'error': 'Username already exists. Please choose another username.'
                          })
        
        if User.objects.filter(email=email).exists():

            return render(request,
                        'public/register.html',
                        {
                            'error': 'Email already registered.'
                        })

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password
        )

        Applicant.objects.create(
            user=user,
            phone='',
            institution='',
            designation='',
            state=''
        )

        return redirect('/login/')

    return render(request,
                  'public/register.html')


# PUBLIC COURSES PAGE

def courses(request):

    all_courses = Course.objects.all()

    return render(request,
                  'public/courses.html',
                  {'courses': all_courses})


# APPLICANT DASHBOARD

@login_required
def applicant_dashboard(request):

    if request.user.is_staff:

        return redirect('/admin-dashboard/')

    return render(request,
                  'applicant/applicant_dashboard.html')


# ADMIN DASHBOARD

@login_required
def admin_dashboard(request):

    if not request.user.is_staff:

        return render(request,'public/access_denied.html')

    return render(request,
                  'admin_panel/admin_dashboard.html')


# CREATE COURSE

@login_required
def create_course(request):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    if request.method == 'POST':

        title = request.POST.get('title')

        course_type = request.POST.get('course_type')

        course_code_map = {
            'Orientation Courses': 'OC',
            'Refresher Courses': 'RC-I',
            'Inter/Multi Disciplinary Refresher Courses': 'RC-II', 
            'Short Term Courses': 'SC', 
            'NEP': 'NEP', 
            }

        course_code = course_code_map.get(course_type, 'GEN')

        mode = request.POST.get('mode')

        start_date = request.POST.get('start_date')

        end_date = request.POST.get('end_date')

        Course.objects.create(

            title=title,

            course_type=course_type,

            course_code=course_code,

            mode=mode,

            start_date=start_date,

            end_date=end_date,

        )

        return redirect('/success/?message=Course created successfully')
    
    return render(
        request,
        'admin_panel/create_course.html'
    )

# MANAGE COURSES

@login_required
def manage_courses(request):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    courses = Course.objects.all()

    return render(request,
                  'admin_panel/manage_courses.html',
                  {'courses': courses})


# EDIT COURSE

@login_required
def edit_course(request, id):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    course = Course.objects.get(id=id)

    if request.method == 'POST':

        course.title = request.POST.get('title')

        course_type = request.POST.get('course_type')

        course_code_map = {
            'Orientation Courses': 'OC',
            'Refresher Courses': 'RC-I',
            'Inter/Multi Disciplinary Refresher Courses': 'RC-II', 
            'Short Term Courses': 'SC', 
            'NEP': 'NEP', 
            }


        course.course_type = request.POST.get('course_type')

        course.course_code = course_code_map.get(course_type, 'GEN')

        course.start_date = request.POST.get('start_date')

        course.end_date = request.POST.get('end_date')

        course.mode = request.POST.get('mode')

        course.seat_status = request.POST.get('seat_status')

        course.is_visible = True if request.POST.get('is_visible') == 'true' else False

        course.save()

        return redirect('/success/?message=Course updated successfully')
    return render(request,
                  'admin_panel/edit_course.html',
                  {'course': course})


# DELETE COURSE

@login_required
def delete_course(request, id):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    course = Course.objects.get(id=id)

    course.delete()

    return redirect('/success/?message=Course deleted successfully')


# APPLY COURSE

@login_required
def apply_course(request, id):

    if request.user.is_staff:

        return redirect('/admin-dashboard/')

    course = Course.objects.get(id=id)

    applicant = Applicant.objects.get(user=request.user)

    already_applied = Application.objects.filter(

        applicant=applicant,
        course=course

    ).exists()

    if already_applied:

        return render(request,
                      'applicant/already_applied.html')

    if request.method == 'POST':

        qualification = request.POST.get('qualification')

        experience = request.POST.get('experience')

        address = request.POST.get('address')

        statement = request.POST.get('statement')

        Application.objects.create(

            applicant=applicant,

            course=course,

            qualification=qualification,

            experience=experience,

            address=address,

            statement=statement

        )

        return redirect('/my-applications/')

    return render(request,
                  'applicant/apply_course.html',
                  {'course': course})

# VIEW APPLICATIONS

@login_required
def view_applications(request):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    applications = Application.objects.all()

    return render(request,
                  'admin_panel/view_applications.html',
                  {'applications': applications})

@login_required
def my_applications(request):

    if request.user.is_staff:

        return redirect('/admin-dashboard/')

    applicant = Applicant.objects.get(user=request.user)

    applications = Application.objects.filter(
        applicant=applicant
    ).order_by('-application_date')

    return render(
        request,
        'applicant/my_applications.html',
        {
            'applications': applications
        }
    )

def info_page(request, title, content):

    return render(request,
                  'public/info_page.html',
                  {
                      'title': title,
                      'content': content
                  })

def objective(request):

    content = """

    <h2>
        Objectives of HRDC-BHU
    </h2>

    <p>
        Objectives of the HRDC-BHU are to enable newly appointed and serving lecturers to:
    </p>

    <ul>

        <li>
            Understand the significance of education in general and higher education in particular in both global and Indian contexts.
        </li>

        <li>
            Understand the relationship between education and economic, socio-economic, and cultural development with special reference to Indian democracy, secularism, and social equity.
        </li>

        <li>
            Acquire and improve teaching skills at the college and university level to achieve the goals of higher education.
        </li>

        <li>
            Remain updated with the latest developments in their respective subjects.
        </li>

        <li>
            Understand the organization and management of colleges and universities and recognize the role of teachers in the overall system.
        </li>

        <li>
            Utilize opportunities for development of personality, creativity, and initiative.
        </li>

        <li>
            Promote computer literacy and the use of ICT in teaching and learning processes.
        </li>

    </ul>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'Objective',
                      'content': content
                  })

def philosophy(request):

    content = """

    <h2>
        Teacher as the Central Element
    </h2>

    <p>
        HRDC's main philosophy is based on the belief that the teacher is central to the entire educational system.
    </p>

    <p>
        While teachers are universally recognized as the pivot of education, adequate opportunities for professional development are often lacking. Therefore, mechanisms must be developed within the framework of a knowledge society to continuously support teacher growth and advancement.
    </p>

    <p>
        Teachers are not merely expected to transmit information. They must also guide students to meet the challenges of life and become responsible citizens along with trained professionals.
    </p>

    <h2>
        Changing Role of Teachers
    </h2>

    <p>
        Earlier, teachers often learned the art of teaching through observation of senior colleagues and mentors. However, modern educational demands require organized and systematic orientation and training programmes for newly appointed teachers.
    </p>

    <p>
        The expansion of higher education in India transformed the system from elitist education to mass-based education. While this expansion is a major achievement, it has also led to challenges in maintaining educational standards.
    </p>

    <p>
        In the present era of rapid knowledge growth and information explosion, teachers carry greater responsibility to continuously improve themselves and remain academically updated.
    </p>

    <h2>
        Educational Technology and Information Technology
    </h2>

    <p>
        Educational technology and Information Technology (IT) have become essential components of modern teaching.
    </p>

    <p>
        UGC has introduced specially designed orientation and refresher courses focused on ICT, e-content development, and digital teaching tools.
    </p>

    <p>
        These programmes aim to:
    </p>

    <ul>

        <li>
            Create internet-literate and computer-literate teachers.
        </li>

        <li>
            Promote the use of software tools in teaching.
        </li>

        <li>
            Encourage e-content development and digital learning methods.
        </li>

        <li>
            Improve technological awareness among teachers irrespective of discipline.
        </li>

    </ul>

    <p>
        Knowledge acquisition is viewed as a two-way process between teachers and students, enabling collective advancement of knowledge.
    </p>

    <h2>
        Continuous Professional Development
    </h2>

    <p>
        Knowledge in every discipline is expanding rapidly. Teachers must continuously update themselves to avoid becoming academically outdated.
    </p>

    <p>
        While many teachers independently work to improve their skills and methodologies, HRDC emphasizes structured and organized orientation programmes to support a large number of teachers at the college and university level.
    </p>

    <h2>
        Orientation Programmes
    </h2>

    <p>
        Orientation programmes for newly appointed lecturers emphasize the role of teachers as agents of socio-economic change and national development.
    </p>

    <p>
        These programmes are flexible rather than rigid and focus on:
    </p>

    <ul>

        <li>
            Developing self-reliance among teachers.
        </li>

        <li>
            Encouraging awareness of social, intellectual, and moral responsibilities.
        </li>

        <li>
            Helping teachers discover their potential.
        </li>

        <li>
            Promoting confidence and awareness in educational responsibilities.
        </li>

    </ul>

    <p>
        The programmes aim to make education relevant, dynamic, and socially meaningful.
    </p>

    <h2>
        National Development and Education
    </h2>

    <p>
        Orientation programmes must create awareness regarding the problems faced by Indian society and establish education as an important solution for national progress.
    </p>

    <p>
        Teachers must understand educational goals in the broader context of constitutional values and national development.
    </p>

    <h2>
        Role of Educational Administrators
    </h2>

    <p>
        HRDC also recognizes that orientation programmes can only succeed when educational administrators understand their significance.
    </p>

    <p>
        Therefore, orientation programmes are also planned for:
    </p>

    <ul>

        <li>
            Heads of Departments
        </li>

        <li>
            Principals
        </li>

        <li>
            Academic officers
        </li>

        <li>
            Educational administrators
        </li>

    </ul>

    <p>
        Such exposure helps administrators actively participate in academic development and encourages better supervision and leadership in higher education.
    </p>

    <h2>
        Exchange of Academic Ideas
    </h2>

    <p>
        HRDC promotes exchange of ideas within the academic and educational environment through interaction among teachers, resource persons, researchers, and students.
    </p>

    <p>
        Lectures by eminent scholars and experts are regularly planned to enhance academic interest and create a vibrant intellectual atmosphere within the university.
    </p>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'Philosophy',
                      'content': content
                  })

def functions(request):

    content = """

    <h2>
        Functions of HRDC-BHU
    </h2>

    <p>
        The Human Resource Development Centre (HRDC-BHU) plans, organizes, implements, monitors, and evaluates orientation programmes for newly appointed college and university lecturers.
    </p>

    <p>
        HRDC also organizes refresher courses for serving teachers along with orientation programmes and workshops for senior administrators, Assistant Registrars, Section Officers, Heads of Departments, Principals, and other academic officers.
    </p>

    <h2>
        Major Functions
    </h2>

    <p>
        Specifically, HRDC-BHU performs the following functions:
    </p>

    <ul>

        <li>
            Formulate orientation programmes according to UGC guidelines.
        </li>

        <li>
            Identify resource persons from various fields of specialization for orientation and refresher courses.
        </li>

        <li>
            Familiarize resource persons with the philosophy and guidelines of the programmes.
        </li>

        <li>
            Maintain a library containing reference and source materials required for courses.
        </li>

        <li>
            Develop specially designed teaching and training materials for effective implementation of programmes.
        </li>

        <li>
            Organize, monitor, and evaluate orientation and refresher courses for teachers.
        </li>

        <li>
            Promote a culture of learning and self-improvement among teachers at the tertiary level.
        </li>

        <li>
            Organize orientation programmes for senior administrators and decision-makers in higher education.
        </li>

        <li>
            Provide opportunities for teachers to exchange experiences and learn from peers.
        </li>

        <li>
            Provide a platform for teachers to remain updated with the latest developments in various disciplines.
        </li>

        <li>
            Encourage teachers to widen their knowledge and pursue research activities.
        </li>

        <li>
            Introduce innovative teaching methods and new developments in higher education.
        </li>

        <li>
            Publish academic materials aimed at enhancing teaching and research capabilities.
        </li>

        <li>
            Conduct capability enhancement programmes for non-academic staff to strengthen the teaching-learning environment.
        </li>

    </ul>

    <h2>
        Academic Planning
    </h2>

    <p>
        Thrust areas and resource persons for each refresher course are decided by the HRDC faculty in consultation with the Course Coordinator.
    </p>

    <h2>
        Meetings with University Officers
    </h2>

    <p>
        HRDC also organizes meetings with university officers and academic administrators in order to:
    </p>

    <ul>

        <li>
            Familiarize administrators with the philosophy and importance of orientation and refresher courses.
        </li>

        <li>
            Encourage them to actively depute teachers for participation in programmes.
        </li>

        <li>
            Help administrators understand their evolving supervisory roles in higher education.
        </li>

        <li>
            Facilitate educational reforms through appropriate modifications in management systems at different levels.
        </li>

    </ul>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'Functions',
                      'content': content
                  })

def responsibility(request):

    content = """

    <h2>
        Responsibilities of HRDC-BHU
    </h2>

    <p>
        The Human Resource Development Centre (HRDC-BHU) caters to the academic and professional needs of teachers and educational administrators from colleges and universities.
    </p>

    <p>
        HRDC-BHU maintains its own official website where all important information, notices, and updates are regularly posted and maintained.
    </p>

    <h2>
        Academic Record Maintenance
    </h2>

    <p>
        In order to make Orientation Programmes (OPs) and Refresher Courses (RCs) more effective, HRDC-BHU maintains systematic records related to:
    </p>

    <ul>

        <li>
            Participants and their academic achievements.
        </li>

        <li>
            Professional growth and teaching capabilities of teachers.
        </li>

        <li>
            Course-wise records of resource persons and participants.
        </li>

        <li>
            Year-wise and subject-wise details of programmes conducted.
        </li>

        <li>
            Reading materials and academic resources produced during courses.
        </li>

    </ul>

    <p>
        Copies of reading materials and publications are also preserved in the HRDC library for academic reference.
    </p>

    <h2>
        Establishment and Growth
    </h2>

    <p>
        At present, there are 66 Human Resource Development Centres functioning across India.
    </p>

    <p>
        The Academic Staff College at Banaras Hindu University was established during the Seventh Five-Year Plan in the year 1987 by the University Grants Commission (UGC).
    </p>

    <h2>
        Directors of HRDC-BHU
    </h2>

    <p>
        The following distinguished academicians have served as Directors of the Centre:
    </p>

    <ul>

        <li>
            Prof. B. B. Dhar
        </li>

        <li>
            Prof. M. S. Srinivasan
        </li>

        <li>
            Prof. V. S. Jaiswal
        </li>

        <li>
            Prof. K. K. Narang
        </li>

        <li>
            Prof. Janardan Singh
        </li>

        <li>
            Prof. Kumar Pankaj
        </li>

        <li>
            Prof. B. B. Bansal
        </li>

        <li>
            Prof. S. N. Upadhyay
        </li>

        <li>
            Prof. Kamal Sheel
        </li>

    </ul>

    <p>
        The present Director of HRDC-BHU is Prof. A. V. Sharma.
    </p>

    <h2>
        Transition to UGC-HRDC
    </h2>

    <p>
        From April 01, 2015, the University Grants Commission officially changed the name of UGC-Academic Staff Colleges to UGC-Human Resource Development Centres (UGC-HRDC).
    </p>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'Responsibility',
                      'content': content
                  })

def isro_edusat(request):

    content = """

    <h2>
        ISRO-EDUSAT Centre
    </h2>

    <p>
        HRDC-BHU has an EDUSAT Centre which regularly receives educational transmissions from organizations such as:
    </p>

    <ul>

        <li>
            CEC-UGC
        </li>

        <li>
            NCERT
        </li>

        <li>
            DST
        </li>

        <li>
            ISRO
        </li>

        <li>
            IGNOU
        </li>

        <li>
            Indian Institute of Remote Sensing (IIRS)
        </li>

    </ul>

    <p>
        The centre facilitates live interaction between experts and participants attending various courses organized by HRDC-BHU.
    </p>

    <p>
        The EDUSAT Centre has been recognized by the Indian Institute of Remote Sensing (IIRS), Dehradun as a University Centre for EDUSAT-based distance learning programmes.
    </p>

    <p>
        Since 2009 onwards, the centre has been regularly organizing programmes related to:
    </p>

    <ul>

        <li>
            Remote Sensing
        </li>

        <li>
            Geographic Information System (GIS)
        </li>

        <li>
            Global Positioning System (GPS)
        </li>

    </ul>

    <h2>
        Computer Laboratory
    </h2>

    <p>
        HRDC-BHU also maintains a well-equipped modern computer laboratory established with the financial assistance of UGC.
    </p>

    <p>
        The laboratory contains 25 Core i3 computer systems to support training, academic activities, and ICT-based learning programmes.
    </p>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'ISRO-EDUSAT Centre',
                      'content': content
                  })

def facilities(request):

    content = """

    <h2>
        Smart Classroom
    </h2>

    <p>
        HRDC-BHU has a modern Smart Classroom facility with a seating capacity of 70 participants.
    </p>

    <h2>
        Hostel and Guest House Facility
    </h2>

    <p>
        HRDC-BHU has its own guest house facility to accommodate a limited number of participants on a first-come-first-served basis.
    </p>

    <p>
        In addition to this, accommodation is also available in other guest houses of Banaras Hindu University on payment basis according to the respective guest house rates.
    </p>

    <p>
        Available guest house facilities include:
    </p>

    <ul>

        <li>
            University Guest House
        </li>

        <li>
            Faculty Guest House
        </li>

        <li>
            IIT Guest House
        </li>

        <li>
            Bhatnagar Guest House
        </li>

        <li>
            Transit House
        </li>

    </ul>

    <h2>
        Multipurpose Hall
    </h2>

    <p>
        UGC-HRDC, BHU has a dedicated multipurpose hall for academic and training activities.
    </p>

    <h2>
        Computer Laboratory
    </h2>

    <p>
        HRDC-BHU maintains a well-established Computer Laboratory equipped with 25 personal computers.
    </p>

    <p>
        Internet facilities are available on each system for participants and faculty members.
    </p>

    <h2>
        Library
    </h2>

    <p>
        UGC-HRDC, BHU has its own library containing approximately 2305 books along with monthly subscribed magazines.
    </p>

    <p>
        Participants are also provided access to:
    </p>

    <ul>

        <li>
            Cyber Library of BHU
        </li>

        <li>
            Central Library of BHU
        </li>

    </ul>

    <h2>
        Feedback System
    </h2>

    <p>
        Participants of every course regularly submit feedback regarding lectures and programme activities.
    </p>

    <p>
        At the end of each programme, participants also submit a comprehensive report of the course.
    </p>

    <h2>
        Best Practices
    </h2>

    <ul>

        <li>
            HRDC-BHU distributes plant saplings instead of bouquets to guests during formal functions.
        </li>

        <li>
            A 30-minute review session is conducted daily in addition to the regular six-hour contact sessions.
        </li>

        <li>
            Paperless programmes are encouraged and implemented.
        </li>

        <li>
            Participants are facilitated to use the university Cyber Library from 8:00 A.M. to 11:00 P.M.
        </li>

    </ul>

    """

    return render(request,
                  'public/info_page.html',
                  {
                      'title': 'Facilities',
                      'content': content
                  })

def programme_courses(request):

    content = """

    <h2>
        Full-Time Programmes (Orientation Programmes and Refresher Courses)
    </h2>

    <ul>

        <li>
            Programmes organized by BHU-HRDC are conducted on a full-time basis and maintain a residential character throughout the duration of the course.
        </li>

        <li>
            Participating lecturers are deputed by their respective colleges and universities for the entire duration of the programme.
        </li>

        <li>
            Teachers attending the programme are treated as being on duty with full pay and allowances by their sponsoring institutions.
        </li>

        <li>
            Participants are selected from institutions across India to promote national integration.
        </li>

        <li>
            A course generally consists of 20 or more participants.
        </li>

        <li>
            If sufficient participants are not available, BHU-HRDC coordinates with nearby HRDCs to exchange participants and ensure optimal enrollment.
        </li>

        <li>
            All courses are organized through BHU-HRDC.
        </li>

        <li>
            While conducting Refresher Courses, active involvement of the concerned department faculty is ensured.
        </li>

        <li>
            Punctuality, regularity, active participation, and purposefulness are strongly emphasized.
        </li>

        <li>
            Successful participants are awarded certificates. BHU-HRDC may withhold certificates for valid reasons.
        </li>

        <li>
            Centres of Excellence and specialized institutions may seek UGC approval to conduct sponsored Refresher Courses based on merit and relevance.
        </li>

        <li>
            E-content development and learning-object technologies are integrated into Orientation Programmes and Refresher Courses, with additional e-courses planned for the future.
        </li>

    </ul>

    <h2>
        Short-Term Courses (2–6 Days)
    </h2>

    <p>
        Apart from Orientation Programmes and Refresher Courses, BHU-HRDC plans to conduct short-term programmes ranging from two to six days.
    </p>

    <p>
        These programmes are primarily intended for the professional development of senior faculty members, including Professors and Associate Professors.
    </p>

    <h2>
        Interaction Programmes (21 Days)
    </h2>

    <p>
        Students pursuing Ph.D. and Post-Doctoral studies from Centres of Advanced Studies (CAS) and Departments of Special Assistance (DSA) may participate in special interaction programmes organized under the Refresher Course scheme.
    </p>

    <p>
        These programmes are generally conducted in the form of workshops and seminars with a duration of approximately three weeks.
    </p>

    <p>
        The primary objective is to facilitate meaningful interaction between research scholars and teachers.
    </p>

    <p>
        Student participants receive TA/DA benefits similar to those provided to Refresher Course participants.
    </p>

    <h2>
        Orientation Programmes for Administrative and Non-Academic Staff
    </h2>

    <p>
        BHU-HRDC also plans to organize orientation programmes, seminars, and workshops for academic planners and administrators, including Group 'A' officers.
    </p>

    <p>
        In addition, professional development programmes of approximately six days are proposed for non-academic Group 'B' and Group 'C' staff with the support of UGC.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Courses',
            'content': content
        }
    )

def curriculum(request):

    content = """

    <h2>
        Course Preparation for Refresher Courses
    </h2>

    <p>
        The curriculum of Refresher Courses is developed in consultation with the concerned academic departments.
    </p>

    <p>
        The course structure includes core subject content, emerging areas of knowledge, practical and laboratory components, computer applications, and recent advancements relevant to the discipline.
    </p>

    <h2>
        Components of Orientation Programmes
    </h2>

    <p>
        Orientation Programmes consist of four major components with a minimum of 144 contact hours spread over 24 working days.
    </p>

    <h3>
        Component A: Society, Environment, Development and Education
    </h3>

    <ul>
        <li>Constitution of India and secularism</li>
        <li>National integration</li>
        <li>Status of women and children</li>
        <li>Inclusive development</li>
        <li>Environmental awareness and biodiversity</li>
        <li>Economic issues and national development</li>
        <li>Urbanization and modernization</li>
        <li>Youth power</li>
        <li>Role and responsibility of teachers</li>
        <li>Value-based education</li>
        <li>Indian culture and identity</li>
        <li>Human rights</li>
        <li>Sustainable development</li>
        <li>Globalization, privatization and liberalization</li>
        <li>Public interest movements</li>
        <li>Aesthetics</li>
    </ul>

    <h3>
        Component B: Philosophy of Education and Pedagogy
    </h3>

    <ul>
        <li>Philosophy of education</li>
        <li>Indian education system and policies</li>
        <li>Economics of education</li>
        <li>Quality assurance and accreditation</li>
        <li>Learner psychology and motivation</li>
        <li>Teaching methods and classroom techniques</li>
        <li>Educational technology and ICT</li>
        <li>Curriculum design</li>
        <li>Evaluation and examination reforms</li>
        <li>Distance and open learning systems</li>
    </ul>

    <h3>
        Component C: Resource Awareness and Knowledge Generation
    </h3>

    <ul>
        <li>Information and Communication Technology</li>
        <li>Information networks and databases</li>
        <li>Libraries and documentation centres</li>
        <li>Museums, laboratories and centres of excellence</li>
        <li>Research projects and publications</li>
        <li>Industry-university linkages</li>
    </ul>

    <h3>
        Component D: Management and Personality Development
    </h3>

    <ul>
        <li>Communication skills</li>
        <li>Scientific temper and critical thinking</li>
        <li>Creativity and innovation</li>
        <li>Leadership and team building</li>
        <li>Administrative skills</li>
        <li>Educational management</li>
        <li>Student guidance and counselling</li>
        <li>Mental health and values</li>
        <li>Career planning and time management</li>
        <li>Teacher effectiveness and accountability</li>
    </ul>

    <h2>
        Information Technology Orientation
    </h2>

    <p>
        HRDC organizes interdisciplinary Refresher Courses in IT Awareness. IT orientation is also incorporated into other Orientation Programmes and Refresher Courses.
    </p>

    <h2>
        Eligibility and Target Group
    </h2>

    <p>
        Teachers working in universities and colleges covered under Section 2(f) of the UGC Act are eligible to participate in Orientation Programmes and Refresher Courses.
    </p>

    <ul>
        <li>Newly appointed lecturers with up to six years of service may attend Orientation Programmes.</li>
        <li>Completion of an Orientation Programme is generally required before attending a Refresher Course.</li>
        <li>A minimum gap of one year is normally maintained between two courses.</li>
        <li>Part-time, ad-hoc, temporary and contract teachers may also participate subject to eligibility conditions.</li>
    </ul>

    <h2>
        Duration of Programmes
    </h2>

    <ul>
        <li>
            Orientation Programme: 4 weeks, 24 working days, 144 contact hours.
        </li>

        <li>
            Refresher Course: 3 weeks, 18 working days, 108 contact hours.
        </li>
    </ul>

    <p>
        Participants who fail to complete the required contact hours may be allowed to make up the deficiency in another programme at their own expense.
    </p>

    <h2>
        Refresher Courses During Teacher Fellowship
    </h2>

    <p>
        Teachers on fellowship may attend Refresher Courses relevant to their research area provided they surrender fellowship living expenses for the course period and submit the required undertaking.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Curriculum',
            'content': content
        }
    )

def course_coordinator(request):

    content = """

    <h2>
        Course Coordinator
    </h2>

    <p>
        The HRDC may appoint, if required, one Coordinator for a Refresher Course.
    </p>

    <p>
        A lump-sum honorarium of Rs. 6,000 is admissible to the Coordinator for the course.
    </p>

    <p>
        In special circumstances, more than one Coordinator may be appointed. In such cases, the honorarium amount is equally shared among the Coordinators.
    </p>

    <p>
        Coordinators are not entitled to receive honorarium for taking classes in the same course for which they are serving as Coordinator.
    </p>

    <h2>
        Orientation Programmes
    </h2>

    <p>
        In the case of Orientation Programmes, the Director may appoint one member of the academic staff as the Coordinator of the programme.
    </p>

    <p>
        Such Coordinators are not entitled to receive any honorarium for coordinating the programme.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Course Coordinator',
            'content': content
        }
    )

def teaching_staff(request):

    return render(
        request,
        'public/teaching_staff.html',
        {
            'title': 'Teaching Staff'
        }
    )

def non_teaching_staff(request):

    return render(
        request,
        'public/non_teaching_staff.html',
        {
            'title': 'Non Teaching Staff'
        }
    )

def activities(request):

    content = """

    <h2>
        Major Activities
    </h2>

    <ol>
        <li>
            National Seminar-cum-Workshop Series on Research Methodologies and Communication Skills (REMECOS), 19th March to 20th March 2007, organised by ASC-BHU.
        </li>

        <li>
            National Conference on Air Pollution and Health Effects (NACAPHE), 28th March to 29th March 2007, organised by ASC-BHU.
        </li>
    </ol>

    <h2>
        Workshops / Seminars Organized in Collaboration with Departments
    </h2>

    <table class="info-table">

        <tr>
            <th>Sl. No.</th>
            <th>Title of the Program</th>
            <th>Date / Duration</th>
            <th>Male</th>
            <th>Female</th>
            <th>Total</th>
        </tr>

        <tr>
            <td>1</td>
            <td>Non-conventional Energy Resources in Dry Tropics (Dept. of Physics, BHU)</td>
            <td>30th April, 1990</td>
            <td>28</td>
            <td>19</td>
            <td>47</td>
        </tr>

        <tr>
            <td>2</td>
            <td>Fractals (Dept. of Physics, BHU)</td>
            <td>2nd May, 1990</td>
            <td>32</td>
            <td>10</td>
            <td>42</td>
        </tr>

        <tr>
            <td>3</td>
            <td>Role of Growth Centres in the Economic Development of Eastern U.P. (Dept. of Economics, BHU)</td>
            <td>14th July, 1990</td>
            <td>28</td>
            <td>15</td>
            <td>43</td>
        </tr>

        <tr>
            <td>4</td>
            <td>Girl Child and Family (Dept. of Sociology, BHU)</td>
            <td>4th August, 1990</td>
            <td>15</td>
            <td>35</td>
            <td>50</td>
        </tr>

        <tr>
            <td>5</td>
            <td>Prof. U. C. Nag Memorial Annual Lecture (Dept. of English, BHU)</td>
            <td>Sept. 9-11, 1993</td>
            <td>30</td>
            <td>35</td>
            <td>65</td>
        </tr>

        <tr>
            <td>6</td>
            <td>New Directions in Indian English Literature (Dept. of English)</td>
            <td>Sept. 8-10, 1994</td>
            <td>25</td>
            <td>30</td>
            <td>55</td>
        </tr>

        <tr>
            <td>7</td>
            <td>University Accounts and Audits (Office of Registrar, BHU)</td>
            <td>December 6-10, 1994</td>
            <td>25</td>
            <td>05</td>
            <td>30</td>
        </tr>

        <tr>
            <td>8</td>
            <td>XVII Dr. U. G. Nag Memorial Lecture (Dept. of English, BHU)</td>
            <td>January 27-28, 1995</td>
            <td>30</td>
            <td>40</td>
            <td>70</td>
        </tr>

        <tr>
            <td>9</td>
            <td>Ecology: Current Problem and Future Perspective (Dept. of Botany, BHU)</td>
            <td>March 10-12, 1995</td>
            <td>40</td>
            <td>25</td>
            <td>65</td>
        </tr>

        <tr>
            <td>10</td>
            <td>Academic Interaction between India and Canada: Retrospect & Prospect (CSDP, Dept. of Political Science, BHU)</td>
            <td>Nov. 19-20, 1995</td>
            <td>52</td>
            <td>26</td>
            <td>78</td>
        </tr>

        <tr>
            <td>11</td>
            <td>DST/PAC Meeting on Plant Sciences (Dept. of Botany, BHU)</td>
            <td>October 12-13, 1995</td>
            <td>15</td>
            <td>05</td>
            <td>20</td>
        </tr>

        <tr>
            <td>12</td>
            <td>Changing Scenario of Plant Sciences (Dept. of Botany, BHU)</td>
            <td>Oct. 13-15, 1995</td>
            <td>60</td>
            <td>60</td>
            <td>120</td>
        </tr>

        <tr>
            <td>13</td>
            <td>Ethnic Harmony in a Multicultural Society: Lessons from Canada & India (CSDP, Dept. of Political Science, BHU)</td>
            <td>February 21, 1996</td>
            <td>72</td>
            <td>15</td>
            <td>87</td>
        </tr>

        <tr>
            <td>14</td>
            <td>Frontiers of Molecular Biology, Molecular Biology Unit, IMS, BHU</td>
            <td>February 7, 1997</td>
            <td>28</td>
            <td>20</td>
            <td>48</td>
        </tr>

        <tr>
            <td>15</td>
            <td>Challenges of Multiculturalism: India and Canada (CSDP, Dept. of Political Science, BHU)</td>
            <td>February 19-20, 1997</td>
            <td>23</td>
            <td>22</td>
            <td>45</td>
        </tr>

        <tr>
            <td>16</td>
            <td>Regionalism: Canadian and Indian Perspectives (CSDP, Dept. of Political Science, BHU)</td>
            <td>February 3-4, 1998</td>
            <td>50</td>
            <td>25</td>
            <td>75</td>
        </tr>

        <tr>
            <td>17</td>
            <td>English Literature of SAARC Countries (Dept. of English, BHU)</td>
            <td>February 23-25, 1998</td>
            <td>60</td>
            <td>28</td>
            <td>88</td>
        </tr>

        <tr>
            <td>18</td>
            <td>Seminar by CSDP, Dept. of Political Science, BHU</td>
            <td>January 20-21, 1999</td>
            <td>40</td>
            <td>35</td>
            <td>75</td>
        </tr>

        <tr>
            <td>19</td>
            <td>Seminar by CSDP, Dept. of Political Science, BHU</td>
            <td>February 5-6, 2000</td>
            <td>35</td>
            <td>30</td>
            <td>55</td>
        </tr>

        <tr>
            <td>20</td>
            <td>Seminar by CSDP, Dept. of Political Science, BHU</td>
            <td>January 27-28, 2001</td>
            <td>45</td>
            <td>35</td>
            <td>80</td>
        </tr>

        <tr>
            <td>21</td>
            <td>Seminar by CSDP, Dept. of Political Science, BHU</td>
            <td>February 10-11, 2001</td>
            <td>30</td>
            <td>40</td>
            <td>70</td>
        </tr>

        <tr>
            <td>22</td>
            <td>Higher Education in India and Canada (CSDP, Dept. of Political Science, BHU)</td>
            <td>January 24-25, 2003</td>
            <td>38</td>
            <td>40</td>
            <td>78</td>
        </tr>

        <tr>
            <td>23</td>
            <td>VIII Group Meeting of AINP on Agricultural Acarology (IAS, BHU)</td>
            <td>March 28-29, 2003</td>
            <td>25</td>
            <td>20</td>
            <td>45</td>
        </tr>

        <tr>
            <td>24</td>
            <td>Winter School on Communication Support for Sustainable Development (Institute of Agricultural Sciences, BHU)</td>
            <td>October 9-29, 2004</td>
            <td>20</td>
            <td>07</td>
            <td>27</td>
        </tr>

        <tr>
            <td>25</td>
            <td>International Seminar on Nissim Ezekiel and Indian Literature in English</td>
            <td>January 18-19, 2005</td>
            <td>90</td>
            <td>50</td>
            <td>140</td>
        </tr>

        <tr>
            <td>26</td>
            <td>Seminar on Indo-Canadian Relations by CSDP, Dept. of Political Science, BHU</td>
            <td>November 21-22, 2005</td>
            <td>20</td>
            <td>35</td>
            <td>55</td>
        </tr>

        <tr>
            <td>27</td>
            <td>National Seminar on Digital Preservation of Manuscripts and Rare Materials by Central Library, BHU</td>
            <td>December 24, 2005</td>
            <td>30</td>
            <td>15</td>
            <td>45</td>
        </tr>

        <tr>
            <td>28</td>
            <td>Workshop by Dept. of English, BHU</td>
            <td>Feb. 7-8, 2006</td>
            <td>35</td>
            <td>10</td>
            <td>45</td>
        </tr>

        <tr>
            <td>29</td>
            <td>International Seminar on W. H. Auden and His Contemporary Relevance by Dept. of English</td>
            <td>September 14-15, 2006</td>
            <td>45</td>
            <td>20</td>
            <td>65</td>
        </tr>

    </table>

    <h2>
        Other Programs Organized
    </h2>

    <table class="info-table">

        <tr>
            <th>Sl. No.</th>
            <th>Title of the Program</th>
            <th>Date / Duration</th>
            <th>Male</th>
            <th>Female</th>
            <th>Total</th>
        </tr>

        <tr>
            <td>1</td>
            <td>Engineering Course (AICTE funded, New Delhi)</td>
            <td>May 8-29, 1995</td>
            <td>20</td>
            <td>04</td>
            <td>24</td>
        </tr>

        <tr>
            <td>2</td>
            <td>RC in Philosophy (ICPR funded, New Delhi)</td>
            <td>Feb. 8-28, 2005</td>
            <td>31</td>
            <td>09</td>
            <td>40</td>
        </tr>

        <tr>
            <td>3</td>
            <td>Workshop on Healthy Aging</td>
            <td>Nov. 14-18, 2005</td>
            <td>40</td>
            <td>15</td>
            <td>55</td>
        </tr>

    </table>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Activities',
            'content': content
        }
    )

def terms_conditions(request):

    content = """

    <h2>
        Eligibility and Terms & Conditions
    </h2>

    <ol>

        <li>
            Teachers working in Universities and Colleges included under Section 2(f) of the UGC Act are eligible to participate in Orientation Programmes and Refresher Courses. Teachers of colleges affiliated to a university for at least two years may also participate. However, such participants will not be entitled to TA/DA and other allowances.
            <br><br>
            For Orientation Programmes, newly appointed Assistant Professors within two years of continuous service and teachers requiring orientation for career advancement are eligible.
            <br><br>
            For Refresher Courses, the eligibility requirement for teachers who have not attended an Orientation Course has been reduced to two years. Normally, a gap of one year should be maintained between two Refresher Courses.
        </li>

        <br>

        <li>
            Every participant must pay a non-refundable registration fee of Rs. 1000 through Demand Draft along with the acceptance letter.
        </li>

        <br>

        <li>
            Limited dormitory-type accommodation is available for outstation participants on a chargeable basis. Participants will receive TA/DA according to UGC norms. Boarding expenses may have to be partially borne by participants.
        </li>

        <br>

        <li>
            Outstation participants travelling from a distance of more than 500 km will be reimbursed single return First Class railway fare through the shortest route on successful completion of the course. Original railway tickets for both onward and return journeys must be submitted. Otherwise, reimbursement will be restricted to Sleeper Class fare.
            <br><br>
            Participants travelling from less than 500 km will be reimbursed Second Class fare only.
        </li>

        <br>

        <li>
            Free boarding and lodging facilities are provided to outstation participants during the programme. No separate DA will be paid for the course period. Local participants are provided lunch, tea and refreshments as per UGC norms.
        </li>

        <br>

        <li>
            Participants must bring two passport-size photographs at the time of registration.
        </li>

        <br>

        <li>
            The course is conducted on a full-time basis, six days a week, including holidays except Sundays. A detailed programme schedule is provided during registration.
        </li>

        <br>

        <li>
            Every participant is required to deliver a seminar or lecture of 15 minutes duration followed by 5 minutes of discussion on a topic related to the course theme.
            <br><br>
            Participants must bring their own presentation materials and submit a typed write-up of the seminar on A4-size paper at the time of registration.
        </li>

        <br>

        <li>
            All participants must complete registration on the inaugural day at the Academic Staff College, Central Library Complex, BHU. Failure to do so will result in cancellation of admission.
        </li>

        <br>

        <li>
            Participants are expected to provide feedback and suggestions regarding lectures, programme components, and fellow participants' seminars.
        </li>

        <br>

        <li>
            Participants must attend all sessions sincerely and regularly. No leave will be granted during the programme.
        </li>

        <br>

        <li>
            Certificates of participation will be awarded only to those participants who successfully complete all programme requirements.
            <br><br>
            In case of any violation or default, the Director of the Academic Staff College, BHU reserves the right to cancel admission, withhold TA/DA payments, or deny issuance of certificates. Decisions in such cases shall be final.
        </li>

    </ol>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Terms and Conditions',
            'content': content
        }
    )

def evaluation_participants(request):

    content = """

    <h2>
        Evaluation of Participants
    </h2>

    <p>
        Whenever required, participants may be assessed by experts, preferably external experts, through multiple-choice objective tests and other evaluation methods conducted during the programme.
    </p>

    <h2>
        Grading System
    </h2>

    <table class="info-table">

        <tr>
            <th>Grade</th>
            <th>Percentage</th>
        </tr>

        <tr>
            <td>A</td>
            <td>75% and above</td>
        </tr>

        <tr>
            <td>B</td>
            <td>60% to less than 75%</td>
        </tr>

        <tr>
            <td>C</td>
            <td>50% to less than 60%</td>
        </tr>

        <tr>
            <td>D</td>
            <td>Below 50%</td>
        </tr>

    </table>

    <p>
        Participants obtaining Grade D are required to repeat the programme after a gap of one year without financial commitment from UGC-HRDC.
    </p>

    <h2>
        Distribution of Marks
    </h2>

    <table class="info-table">

        <tr>
            <th>Component</th>
            <th>Marks</th>
        </tr>

        <tr>
            <td>Multiple-choice Objective Tests</td>
            <td>30</td>
        </tr>

        <tr>
            <td>Seminars / Participant Presentation</td>
            <td>15</td>
        </tr>

        <tr>
            <td>Project / Survey / Other Activities</td>
            <td>20</td>
        </tr>

        <tr>
            <td>Micro-Teaching / Participation</td>
            <td>10</td>
        </tr>

        <tr>
            <td>Holistic Response (Punctuality, Regularity, Initiative, Conduct, Responsiveness, etc.)</td>
            <td>25</td>
        </tr>

        <tr>
            <th>Total</th>
            <th>100</th>
        </tr>

    </table>

    <p>
        The distribution of marks may be modified to suit specific programme requirements.
    </p>

    <p>
        Certificates issued to participants are valid only when supported by the grade obtained through the evaluation process.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Evaluation of Participants',
            'content': content
        }
    )

def resource_persons(request):

    content = """

    <h2>
        Resource Persons
    </h2>

    <p>
        Resource persons play an important role in the successful conduct of Orientation Programmes, Refresher Courses, Workshops, and other academic activities organized by HRDC-BHU.
    </p>

    <h2>
        Travel Allowance and Daily Allowance
    </h2>

    <p>
        Outstation resource persons are paid Travel Allowance (TA) and Daily Allowance (DA) according to the prescribed norms.
    </p>

    <h2>
        Honorarium
    </h2>

    <p>
        Honorarium may be paid to both local and outstation resource persons at the rate of:
    </p>

    <ul>

        <li>
            Rs. 1,500 per session of 90 minutes.
        </li>

        <li>
            Maximum Rs. 3,000 per day.
        </li>

        <li>
            Maximum Rs. 6,000 per course.
        </li>

    </ul>

    <h2>
        Outstation Resource Persons
    </h2>

    <p>
        Outstation resource persons are normally invited only once during a programme.
    </p>

    <h2>
        Local Resource Persons
    </h2>

    <p>
        Local resource persons may be reimbursed actual conveyance charges up to Rs. 500 each way.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Resource Persons',
            'content': content
        }
    )

def contact_us(request):

    if request.method == 'POST':

        name = request.POST.get('name')

        email = request.POST.get('email')

        subject = request.POST.get('subject')

        message = request.POST.get('message')

        ContactMessage.objects.create(
            name=name,
            email=email,
            subject=subject,
            message=message
        )

        return render(
            request,
            'public/contact_us.html',
            {'success': True}
        )

    return render(
        request,
        'public/contact_us.html'
    )

def gallery(request):

    content = """

    <h2>
        Gallery
    </h2>

    <p style="text-align:center;
              font-size:20px;
              padding:40px 0;">
        Photos will appear here.
    </p>

    """

    return render(
        request,
        'public/info_page.html',
        {
            'title': 'Gallery',
            'content': content
        }
    )

@login_required
def view_messages(request):

    if not request.user.is_staff:

        return redirect('/dashboard/')

    messages = ContactMessage.objects.all().order_by('-created_at')

    return render(
        request,
        'admin_panel/view_messages.html',
        {
            'messages': messages
        }
    )

@login_required
def apply_online(request):

    if request.user.is_staff:
        return redirect('/admin-dashboard/')

    applicant = Applicant.objects.get(user=request.user)

    courses = Course.objects.filter(
        is_visible=True,
        seat_status='Vacant'
    ).order_by('course_code', 'start_date')

    if request.method == 'POST':

        course = Course.objects.get(id=request.POST.get('course'))

        if Application.objects.filter(applicant=applicant, course=course).exists():

            return render(request,
                          'applicant/apply_online.html',
                          {
                              'courses': courses,
                              'error': 'You have already applied for this course.'
                          })

        application = Application.objects.create(

            applicant=applicant,
            course=course,

            title=request.POST.get('title'),
            full_name=request.POST.get('full_name'),
            category=request.POST.get('category'),
            is_pwd=True if request.POST.get('is_pwd') == 'YES' else False,
            aadhaar_number=request.POST.get('aadhaar_number'),
            photograph=request.FILES.get('photograph'),
            date_of_birth=request.POST.get('date_of_birth'),
            application_number=generate_application_number(),
            applicant_type=request.POST.get('applicant_type'),
            present_designation=request.POST.get('present_designation'),
            appointment_date=request.POST.get('appointment_date') or None,
            scale_of_pay=request.POST.get('scale_of_pay'),
            nature_of_appointment=request.POST.get('nature_of_appointment'),

            official_address=request.POST.get('official_address'),
            official_state=request.POST.get('official_state'),
            official_pin=request.POST.get('official_pin'),

            mailing_address=request.POST.get('mailing_address'),
            mailing_state=request.POST.get('mailing_state'),
            mailing_pin=request.POST.get('mailing_pin'),

            mobile_number=request.POST.get('mobile_number'),
            alternate_number=request.POST.get('alternate_number'),
            email=request.POST.get('email'),

            highest_qualification=request.POST.get('highest_qualification'),
            qualification_status=request.POST.get('qualification_status'),
            subject_specialization=request.POST.get('subject_specialization'),

            accommodation_required=True if request.POST.get('accommodation_required') == 'YES' else False,

            status='Submitted'
        )

        for i in range(1, 5):

            course_name = request.POST.get(f'earlier_course_name_{i}')
            month_year = request.POST.get(f'completion_month_year_{i}')
            center_name = request.POST.get(f'center_name_{i}')

            if course_name or month_year or center_name:

                EarlierAttendedCourse.objects.create(
                    application=application,
                    course_name=course_name,
                    completion_month_year=month_year,
                    center_name=center_name
                )
            generate_application_files(application)

        return redirect(
            f'/success/?message=Application submitted successfully&application_id={application.id}')

    return render(request,
                  'applicant/apply_online.html',
                  {'courses': courses})


@login_required
def reupload_form(request):

    if request.user.is_staff:
        return redirect('/admin-dashboard/')

    applicant = Applicant.objects.get(user=request.user)

    applications = Application.objects.filter(
        applicant=applicant
    ).order_by('-application_date')

    if request.method == 'POST':

        application_id = request.POST.get('application_id')

        application = Application.objects.get(
            id=application_id,
            applicant=applicant
        )

        if application.signed_form:

            return redirect(
                '/success/?message=Signed form has already been uploaded for this application.'
            )

        application.signed_form = request.FILES.get('signed_form')

        application.status = 'Signed Form Uploaded'

        application.save()

        return redirect(
            '/success/?message=Signed form uploaded successfully.'
        )

    return render(
        request,
        'applicant/reupload_form.html',
        {
            'applications': applications
        }
    )


@login_required
def application_status(request):

    if request.user.is_staff:
        return redirect('/admin-dashboard/')

    applicant = Applicant.objects.get(user=request.user)

    applications = Application.objects.filter(
        applicant=applicant
    ).order_by('-application_date')

    search = request.GET.get('search')

    if search:

        applications = applications.filter(
            application_number=search
        )

    return render(
        request,
        'applicant/application_status.html',
        {
            'applications': applications,
            'search': search,
        }
    )


@login_required
def certificates(request):

    return render(request, 'applicant/certificates.html')

def success_page(request):

    message = request.GET.get(
        'message',
        'Operation completed successfully.'
    )

    application_id = request.GET.get('application_id')

    application = None

    if application_id:

        application = Application.objects.get(id=application_id)

    return render(
        request,
        'public/success.html',
        {
            'message': message,
            'application': application
        }
    )

@login_required
def update_course_status(request, id):

    if not request.user.is_staff:
        return redirect('/dashboard/')

    course = Course.objects.get(id=id)

    if request.method == 'POST':

        course.seat_status = request.POST.get('seat_status')

        course.is_visible = True if request.POST.get('is_visible') == 'true' else False

        course.save()

        return redirect('/success/?message=Course status updated successfully')
    
@login_required
def mark_course_full(request, id):

    course = Course.objects.get(id=id)

    course.seat_status = 'Full'

    course.save()

    return redirect('/manage-courses/')

@login_required
def mark_course_vacant(request, id):

    course = Course.objects.get(id=id)

    course.seat_status = 'Vacant'

    course.save()

    return redirect('/manage-courses/')

@login_required
def hide_course(request, id):

    course = Course.objects.get(id=id)

    course.is_visible = False

    course.save()

    return redirect('/manage-courses/')

@login_required
def show_course(request, id):

    course = Course.objects.get(id=id)

    course.is_visible = True

    course.save()

    return redirect('/manage-courses/')

def working_days(start_date, end_date):

    count = 0

    current = start_date

    while current <= end_date:

        if current.weekday() != 6:
            count += 1

        current += timedelta(days=1)

    return count

@login_required
def generate_schedule_pdf(request):

    if not request.user.is_staff:
        return redirect('/dashboard/')

    order = {
        'OC': 1,
        'RC-I': 2,
        'RC-II': 3,
        'SC': 4,
        'NEP': 5,
    }

    course_type_names = {
        'OC': 'ORIENTATION COURSES',
        'RC-I': 'REFRESHER COURSES',
        'RC-II': 'INTER/MULTI DISCIPLINARY REFRESHER COURSES',
        'SC': 'SHORT TERM COURSES',
        'NEP': 'NEP COURSES',
    }

    courses = Course.objects.filter(
        is_visible=True
    )

    courses = sorted(
        courses,
        key=lambda course: (
            order.get(course.course_code, 99),
            course.start_date
        )
    )

    schedule_dir = os.path.join(settings.MEDIA_ROOT, 'course_schedules')
    os.makedirs(schedule_dir, exist_ok=True)

    file_path = os.path.join(schedule_dir, 'tentative_schedule.pdf')

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=20,
        leftMargin=20,
        topMargin=25,
        bottomMargin=25
    )

    styles = getSampleStyleSheet()

    title_style = styles['Title']
    title_style.alignment = 1

    normal_style = styles['BodyText']
    normal_style.fontSize = 8
    normal_style.leading = 10

    elements = []

    elements.append(Paragraph(
        "<b>UGC - HUMAN RESOURCE DEVELOPMENT CENTRE</b>",
        title_style
    ))

    elements.append(Paragraph(
        "<b>BANARAS HINDU UNIVERSITY, VARANASI - 221005</b>",
        title_style
    ))

    elements.append(Spacer(1, 18))

    elements.append(Paragraph(
        "<b>TENTATIVE SCHEDULE</b>",
        title_style
    ))

    elements.append(Spacer(1, 18))

    data = [
        [
            'S.No.',
            'Course Name',
            'Dates',
            'Working Days',
            'Mode'
        ]
    ]

    serial = 1
    current_type = None
    heading_rows = []

    for course in courses:

        if current_type != course.course_code:

            current_type = course.course_code

            heading_rows.append(len(data))

            data.append([
                course_type_names.get(course.course_code, course.course_code),
                '',
                '',
                '',
                ''
            ])

        date_text = (
            f"{course.start_date.strftime('%d-%m-%Y')} "
            f"to {course.end_date.strftime('%d-%m-%Y')}"
        )

        data.append([
            str(serial),
            Paragraph(
                f"<b>{course.course_code} :</b> {course.title}",
                normal_style
            ),
            date_text,
            str(working_days(course.start_date, course.end_date)),
            Paragraph(
                f"<b>{course.mode}</b>",
                normal_style
            )
        ])

        serial += 1

    table = Table(
        data,
        colWidths=[35, 310, 100, 60, 50],
        repeatRows=1
    )

    style_commands = [

        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),

    ]

    for row_index in heading_rows:

        style_commands.extend([

            ('SPAN', (0, row_index), (-1, row_index)),
            ('BACKGROUND', (0, row_index), (-1, row_index), colors.whitesmoke),
            ('FONTNAME', (0, row_index), (-1, row_index), 'Helvetica-Bold'),
            ('FONTSIZE', (0, row_index), (-1, row_index), 9),
            ('ALIGN', (0, row_index), (-1, row_index), 'CENTER'),
            ('TEXTCOLOR', (0, row_index), (-1, row_index), colors.black),

        ])

    table.setStyle(TableStyle(style_commands))

    elements.append(table)

    doc.build(elements)

    schedule, created = CourseSchedule.objects.get_or_create(id=1)

    with open(file_path, 'rb') as pdf:
        schedule.pdf_file.save(
            'tentative_schedule.pdf',
            File(pdf),
            save=True
        )

    return redirect('/success/?message=Tentative schedule PDF generated successfully')

@login_required
def application_instructions(request):

    if request.user.is_staff:
        return redirect('/admin-dashboard/')

    return render(
        request,
        'applicant/application_instructions.html'
    )

def generate_application_number():

    today = datetime.now().strftime('%y%m%d')

    count = Application.objects.count() + 1

    return f"BH{today}{count:04d}"


def replace_text_in_paragraphs(paragraphs, replacements):

    for paragraph in paragraphs:

        full_text = paragraph.text

        for key, value in replacements.items():

            if key in full_text:

                full_text = full_text.replace(key, str(value))

                for run in paragraph.runs:
                    run.text = ""

                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    paragraph.runs[0].bold = True
                else:
                    run = paragraph.add_run(full_text)
                    run.bold = True


def replace_text_in_tables(tables, replacements):

    for table in tables:

        for row in table.rows:

            for cell in row.cells:

                replace_text_in_paragraphs(
                    cell.paragraphs,
                    replacements
                )

                replace_text_in_tables(
                    cell.tables,
                    replacements
                )


def replace_text_in_doc(doc, replacements):

    replace_text_in_paragraphs(
        doc.paragraphs,
        replacements
    )

    replace_text_in_tables(
        doc.tables,
        replacements
    )

    for section in doc.sections:

        replace_text_in_paragraphs(
            section.header.paragraphs,
            replacements
        )

        replace_text_in_tables(
            section.header.tables,
            replacements
        )

        replace_text_in_paragraphs(
            section.footer.paragraphs,
            replacements
        )

        replace_text_in_tables(
            section.footer.tables,
            replacements
        )

        for section in doc.sections:

            print("HEADER TEXT:")

            for paragraph in section.header.paragraphs:
                print(paragraph.text)

            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        print(cell.text)

def generate_application_docx(application):

    template_path = os.path.join(
        settings.BASE_DIR,
        'core',
        'templates_docs',
        'Print_Application.docx'
    )

    output_dir = os.path.join(
        settings.MEDIA_ROOT,
        'application_uploads',
        'generated_docs'
    )

    os.makedirs(output_dir, exist_ok=True)

    output_filename = f"{application.application_number}.docx"

    output_path = os.path.join(
        output_dir,
        output_filename
    )

    course_name = (
        f"{application.course.course_code}: {application.course.title} "
        f"({application.course.mode} Mode, "
        f"{application.course.start_date.strftime('%B %d')} - "
        f"{application.course.end_date.strftime('%B %d, %Y')})"
    )

    earlier_courses = list(application.earlier_courses.all())

    replacements = {
        '{{PRINT_DATETIME}}': timezone.now().strftime('%d-%m-%Y %I:%M %p'),
        '{{APPLICATION_NUMBER}}': application.application_number or '',
        '{{COURSE_NAME}}': course_name,
        '{{FULL_NAME}}': f"{application.title} {application.full_name}",
        '{{AADHAAR}}': application.aadhaar_number or '',
        '{{DOB}}': application.date_of_birth.strftime('%d-%m-%Y') if application.date_of_birth else '',
        '{{PRESENT_DESIG}}': application.present_designation or '',
        '{{APP_DATE}}': application.appointment_date.strftime('%d-%m-%Y') if application.appointment_date else '',
        '{{PAY_SCALE}}': application.scale_of_pay or '',
        '{{CATEGORY}}': application.category or '',
        '{{PWD}}': 'YES' if application.is_pwd else 'NO',
        '{{INST_ADDR}}': application.official_address or '',
        '{{STATE1}}': application.official_state or '',
        '{{PIN1}}': application.official_pin or '',
        '{{MAIL_ADDR}}': application.mailing_address or '',
        '{{STATE2}}': application.mailing_state or '',
        '{{PIN2}}': application.mailing_pin or '',
        '{{MOBILE}}': application.mobile_number or '',
        '{{ALT}}': application.alternate_number or '',
        '{{EMAIL}}': application.email or '',
        '{{H_QUAL}}': application.highest_qualification or '',
        '{{QUAL_S}}': application.qualification_status or '',
        '{{SUBJECT}}': application.subject_specialization or '',
        '{{ACC}}': 'YES' if application.accommodation_required else 'NO',
    }

    for i in range(1, 5):

        if i <= len(earlier_courses):

            earlier = earlier_courses[i - 1]

            replacements[f'{{{{COURSE{i}}}}}'] = earlier.course_name or ''
            replacements[f'{{{{DATE{i}}}}}'] = earlier.completion_month_year or ''
            replacements[f'{{{{CENTER{i}}}}}'] = earlier.center_name or ''

        else:

            replacements[f'{{{{COURSE{i}}}}}'] = ''
            replacements[f'{{{{DATE{i}}}}}'] = ''
            replacements[f'{{{{CENTER{i}}}}}'] = ''

    doc = Document(template_path)

    replace_text_in_doc(doc, replacements)

    doc.save(output_path)

    return output_path

def generate_application_files(application):

    template_path = os.path.join(
        settings.BASE_DIR,
        'core',
        'templates_docs',
        'Print_Application.docx'
    )

    output_dir = os.path.join(
        settings.MEDIA_ROOT,
        'application_uploads',
        'generated_docs'
    )

    pdf_dir = os.path.join(
        settings.MEDIA_ROOT,
        'application_uploads',
        'generated_pdfs'
    )

    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(pdf_dir, exist_ok=True)

    docx_name = f"{application.application_number}.docx"
    pdf_name = f"{application.application_number}.pdf"

    docx_path = os.path.join(output_dir, docx_name)
    pdf_path = os.path.join(pdf_dir, pdf_name)

    course_name = (
        f"{application.course.course_code}: {application.course.title} "
        f"({application.course.mode} Mode, "
        f"{application.course.start_date.strftime('%B %d')} - "
        f"{application.course.end_date.strftime('%B %d, %Y')})"
    )

    earlier_courses = list(application.earlier_courses.all())

    replacements = {
        '{{PRINT_DATETIME}}': timezone.now().strftime('%d-%m-%Y %I:%M %p'),
        '{{APPLICATION_NUMBER}}': application.application_number or '',
        '{{COURSE_NAME}}': course_name,
        '{{FULL_NAME}}': f"{application.title} {application.full_name}",
        '{{AADHAAR}}': application.aadhaar_number or '',
        '{{DOB}}': str(application.date_of_birth) if application.date_of_birth else '',
        '{{PRESENT_DESIG}}': application.present_designation or '',
        '{{APP_DATE}}': str(application.appointment_date) if application.appointment_date else '',
        '{{PAY_SCALE}}': application.scale_of_pay or '',
        '{{CATEGORY}}': application.category or '',
        '{{PWD}}': 'YES' if application.is_pwd else 'NO',
        '{{INST_ADDR}}': application.official_address or '',
        '{{STATE1}}': application.official_state or '',
        '{{PIN1}}': application.official_pin or '',
        '{{MAIL_ADDR}}': application.mailing_address or '',
        '{{STATE2}}': application.mailing_state or '',
        '{{PIN2}}': application.mailing_pin or '',
        '{{MOBILE}}': application.mobile_number or '',
        '{{ALT}}': application.alternate_number or '',
        '{{EMAIL}}': application.email or '',
        '{{H_QUAL}}': application.highest_qualification or '',
        '{{QUAL_S}}': application.qualification_status or '',
        '{{SUBJECT}}': application.subject_specialization or '',
        '{{ACC}}': 'YES' if application.accommodation_required else 'NO',
    }

    for i in range(1, 5):
        if i <= len(earlier_courses):
            earlier = earlier_courses[i - 1]
            replacements[f'{{{{COURSE{i}}}}}'] = earlier.course_name or ''
            replacements[f'{{{{DATE{i}}}}}'] = earlier.completion_month_year or ''
            replacements[f'{{{{CENTER{i}}}}}'] = earlier.center_name or ''
        else:
            replacements[f'{{{{COURSE{i}}}}}'] = ''
            replacements[f'{{{{DATE{i}}}}}'] = ''
            replacements[f'{{{{CENTER{i}}}}}'] = ''

    doc = Document(template_path)

    replace_text_in_doc(doc, replacements)
    replace_photo_placeholder(doc, application)

    doc.save(docx_path)

    application.generated_application_docx.save(
        docx_name,
        File(open(docx_path, 'rb')),
        save=False
    )

    subprocess.run(
        [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            pdf_dir,
            docx_path
        ],
        check=True
    )

    application.generated_application_pdf.save(
        pdf_name,
        File(open(pdf_path, 'rb')),
        save=True
    )

def replace_photo_placeholder(doc, application):

    if not application.photograph:
        return

    photo_path = application.photograph.path

    for paragraph in doc.paragraphs:

        if '{{PHOTO}}' in paragraph.text:

            paragraph.clear()

            run = paragraph.add_run()

            run.add_picture(photo_path, width=Inches(1.2), height=Inches(1.4))

            return

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    if '{{PHOTO}}' in paragraph.text:

                        paragraph.clear()

                        run = paragraph.add_run()

                        run.add_picture(photo_path, width=Inches(1.2), height=Inches(1.4))

                        return
                    
